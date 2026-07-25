import os
import json
import logging
from typing import Any, Dict, List, Tuple
import pypdf
import docx
from core.base_agent import BaseAgent

class ContextManagerAgent(BaseAgent):
    """
    Агент управления контекстом (Context Manager Agent).
    Читает входные файлы в форматах PDF, DOCX, TXT, MD, JSON из папки input/,
    нормализует извлеченную информацию с помощью LLM, выявляет конфликты,
    валидирует результат по knowledge.schema.json и записывает итоговый knowledge.json.
    """
    def __init__(self):
        super().__init__("context_manager")
        # Настройка конкретного лог-файла logs/context.log, как указано в требованиях
        self.logger = self._setup_specific_context_logger()

    def _setup_specific_context_logger(self) -> logging.Logger:
        """Настраивает отдельный лог-файл logs/context.log для Context Manager."""
        log_dir = self.paths.get("paths", {}).get("logs_dir", "logs")
        if not os.path.exists(log_dir):
            os.makedirs(log_dir, exist_ok=True)
            
        logger = logging.getLogger("context_manager_specific")
        logger.setLevel(logging.DEBUG if self.settings.get("system", {}).get("debug", False) else logging.INFO)
        
        # Удаляем существующие хэндлеры
        logger.handlers = []
        
        log_file = os.path.join(log_dir, "context.log")
        file_handler = logging.FileHandler(log_file, encoding="utf-8")
        formatter = logging.Formatter(
            "[%(asctime)s] [%(levelname)s] [ContextManager]: %(message)s",
            datefmt="%Y-%m-%d %H:%M:%S"
        )
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)
        
        # Вывод в консоль для отладки
        console_handler = logging.StreamHandler()
        console_handler.setFormatter(formatter)
        logger.addHandler(console_handler)
        
        return logger

    def parse_txt_md(self, file_path: str) -> str:
        """Чтение текстовых файлов Markdown и TXT."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            self.logger.info(f"Успешно прочитан файл {os.path.basename(file_path)} ({len(content)} символов)")
            return content
        except Exception as e:
            self.logger.error(f"Ошибка чтения текстового файла {file_path}: {e}")
            return ""

    def parse_pdf(self, file_path: str) -> str:
        """Извлечение текста из PDF файлов."""
        try:
            text = ""
            reader = pypdf.PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            self.logger.info(f"Успешно распарсен PDF {os.path.basename(file_path)} ({len(text)} символов, {len(reader.pages)} страниц)")
            return text
        except Exception as e:
            self.logger.error(f"Ошибка парсинга PDF {file_path}: {e}")
            return ""

    def parse_docx(self, file_path: str) -> str:
        """Извлечение текста из файлов Word (DOCX)."""
        try:
            text = ""
            doc = docx.Document(file_path)
            # Извлечение из абзацев
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            # Извлечение из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text:
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            self.logger.info(f"Успешно распарсен DOCX {os.path.basename(file_path)} ({len(text)} символов)")
            return text
        except Exception as e:
            self.logger.error(f"Ошибка парсинга DOCX {file_path}: {e}")
            return ""

    def parse_json(self, file_path: str) -> Dict[str, Any]:
        """Парсинг структурированных файлов JSON."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            self.logger.info(f"Успешно прочитан JSON {os.path.basename(file_path)}")
            return data
        except Exception as e:
            self.logger.error(f"Ошибка парсинга JSON {file_path}: {e}")
            return {}

    def scan_input_directory(self) -> Tuple[str, Dict[str, Any]]:
        """
        Сканирует директорию input/ на наличие поддерживаемых файлов.
        Копирует оригиналы в storage/raw/ и агрегирует JSON данные.
        """
        input_dir = self.paths.get("paths", {}).get("input_dir", "input")
        if not os.path.exists(input_dir):
            input_dir = os.path.join("c:/projects/portfolio_factory", input_dir)
            
        if not os.path.exists(input_dir):
            self.logger.warning(f"Директория входных данных {input_dir} не найдена. Создаем её.")
            os.makedirs(input_dir, exist_ok=True)
            return "", {}

        raw_texts = []
        structured_json_data = {
            "raw_resume_data": {},
            "skills": [],
            "projects": [],
            "technologies": [],
            "repositories": [],
            "extracted_facts": [],
            "metrics": [],
            "lessons_learned": [],
            "decisions": [],
            "evidence": []
        }

        all_files = []
        for root, _, filenames in os.walk(input_dir):
            for file in filenames:
                if not file.startswith("~$"):
                    all_files.append(os.path.join(root, file))
        
        self.logger.info(f"Найдено файлов для сканирования: {len(all_files)}")

        # Создаем директорию storage/raw
        storage_raw = self.paths.get("paths", {}).get("storage_raw_dir", "storage/raw")
        if not os.path.isabs(storage_raw):
            storage_raw = os.path.join("c:/projects/portfolio_factory", storage_raw)
        
        storage_raw_media = os.path.join(storage_raw, "media")
        os.makedirs(storage_raw, exist_ok=True)
        os.makedirs(storage_raw_media, exist_ok=True)
        
        import shutil
        import hashlib
        
        for path in all_files:
            file = os.path.basename(path)
            ext = os.path.splitext(file)[1].lower()

            # Если это картинка, копируем в storage/raw/media и создаем evidence
            if ext in [".png", ".jpg", ".jpeg", ".gif", ".svg", ".webp", ".mp4"]:
                dest = os.path.join(storage_raw_media, file)
                try:
                    shutil.copy2(path, dest)
                    
                    evidence_id = f"ev_media_{hashlib.md5(file.encode()).hexdigest()[:8]}"
                    structured_json_data["evidence"].append({
                        "id": evidence_id,
                        "type": "screenshot" if ext != ".mp4" else "video",
                        "source": "manual_upload",
                        "location": os.path.relpath(dest, "c:/projects/portfolio_factory").replace("\\", "/"),
                        "description": f"Медиа-файл: {file}",
                        "confidence": 1.0
                    })
                    self.logger.info(f"Медиа-файл {file} скопирован и добавлен в Evidence.")
                except Exception as e:
                    self.logger.error(f"Не удалось скопировать медиа {file}: {e}")
                continue

            # Копируем текстовые/pdf файлы в storage/raw
            try:
                shutil.copy2(path, os.path.join(storage_raw, file))
            except Exception as e:
                self.logger.error(f"Не удалось скопировать {file} в storage/raw: {e}")

            if ext in [".txt", ".md"] or ext == "":
                content = self.parse_txt_md(path)
                if content:
                    raw_texts.append(f"--- НАЧАЛО ФАЙЛА {file} ---\n{content}\n--- КОНЕЦ ФАЙЛА {file} ---\n")
            elif ext == ".pdf":
                content = self.parse_pdf(path)
                if content:
                    raw_texts.append(f"--- НАЧАЛО PDF ФАЙЛА {file} ---\n{content}\n--- КОНЕЦ PDF ФАЙЛА {file} ---\n")
            elif ext == ".docx":
                content = self.parse_docx(path)
                if content:
                    raw_texts.append(f"--- НАЧАЛО DOCX ФАЙЛА {file} ---\n{content}\n--- КОНЕЦ DOCX ФАЙЛА {file} ---\n")
            elif ext == ".json":
                data = self.parse_json(path)
                if data:
                    self.logger.info(f"Слияние структурированных данных из JSON: {file}")
                    structured_json_data = self.merge_knowledge_data(structured_json_data, data)
            else:
                self.logger.warning(f"Пропуск файла с неподдерживаемым расширением: {file}")

        combined_text = "\n".join(raw_texts)
        return combined_text, structured_json_data

    def merge_knowledge_data(self, base: Dict[str, Any], new_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Объединяет два объекта структуры базы знаний.
        Решает конфликты для одиночных значений и агрегирует списки.
        """
        merged = {
            "raw_resume_data": base.get("raw_resume_data", {}).copy(),
            "skills": base.get("skills", []).copy(),
            "projects": base.get("projects", []).copy(),
            "technologies": base.get("technologies", []).copy(),
            "repositories": base.get("repositories", []).copy(),
            "extracted_facts": base.get("extracted_facts", []).copy(),
            "metrics": base.get("metrics", []).copy(),
            "lessons_learned": base.get("lessons_learned", []).copy(),
            "decisions": base.get("decisions", []).copy(),
            "evidence": base.get("evidence", []).copy()
        }

        # Миграция плоского списка skills из старого формата base.raw_resume_data.skills
        legacy_base_skills = base.get("raw_resume_data", {}).get("skills", [])
        for skill in legacy_base_skills:
            if isinstance(skill, str):
                skill_id = skill.lower().replace(" ", "_").replace(".", "_")
                if not any(s.get("id") == skill_id for s in merged["skills"]):
                    merged["skills"].append({
                        "id": skill_id,
                        "name": skill,
                        "category": "Uncategorized",
                        "confidence": 0.5
                    })

        # 1. Слияние raw_resume_data
        new_resume = new_data.get("raw_resume_data", {})
        if new_resume:
            base_resume = merged["raw_resume_data"]
            # Слияние personal_info
            base_info = base_resume.setdefault("personal_info", {})
            new_info = new_resume.get("personal_info", {})
            for k, val in new_info.items():
                if val:
                    curr_val = base_info.get(k)
                    if curr_val and curr_val != val and not curr_val.startswith("CONFLICT:"):
                        conflict_str = f"CONFLICT: {curr_val} | {val}"
                        self.logger.warning(f"Конфликт в поле {k}: '{curr_val}' vs '{val}'. Значение сохранено как конфликт.")
                        base_info[k] = conflict_str
                    elif not curr_val:
                        base_info[k] = val

            # Слияние summary
            base_summary = base_resume.get("summary")
            new_summary = new_resume.get("summary")
            if new_summary:
                if base_summary and base_summary != new_summary and not base_summary.startswith("CONFLICT:"):
                    conflict_str = f"CONFLICT: {base_summary} | {new_summary}"
                    self.logger.warning("Конфликт в поле summary. Значение сохранено как конфликт.")
                    base_resume["summary"] = conflict_str
                elif not base_summary:
                    base_resume["summary"] = new_summary

            # Слияние списков (skills, education, work_experience)
            for list_field in ["skills", "education", "work_experience"]:
                base_list = base_resume.setdefault(list_field, [])
                new_list = new_resume.get(list_field, [])
                if new_list:
                    for item in new_list:
                        if item not in base_list:
                            base_list.append(item)

        # 1.5. Слияние skills (и миграция из raw_resume_data.skills)
        base_skills = merged["skills"]
        new_skills = new_data.get("skills", [])
        for skill in new_skills:
            if isinstance(skill, dict) and "id" in skill:
                dup = next((s for s in base_skills if s.get("id") == skill["id"]), None)
                if dup:
                    for k, v in skill.items():
                        if v and not dup.get(k):
                            dup[k] = v
                else:
                    base_skills.append(skill)
            elif isinstance(skill, str):
                skill_id = skill.lower().replace(" ", "_").replace(".", "_")
                if not any(s.get("id") == skill_id for s in base_skills):
                    base_skills.append({
                        "id": skill_id,
                        "name": skill,
                        "category": "Uncategorized",
                        "confidence": 0.5
                    })

        if new_resume:
            legacy_skills = new_resume.get("skills", [])
            for skill in legacy_skills:
                if isinstance(skill, str):
                    skill_id = skill.lower().replace(" ", "_").replace(".", "_")
                    if not any(s.get("id") == skill_id for s in base_skills):
                        base_skills.append({
                            "id": skill_id,
                            "name": skill,
                            "category": "Uncategorized",
                            "confidence": 0.5
                        })

        # 2. Слияние projects
        new_projects = new_data.get("projects", [])
        base_projects = merged["projects"]
        for proj in new_projects:
            proj_id = proj.get("id")
            dup = next((p for p in base_projects if p.get("id") == proj_id), None)
            if dup:
                for k, v in proj.items():
                    if v and not dup.get(k):
                        dup[k] = v
            else:
                base_projects.append(proj)

        # 3. Слияние technologies
        new_techs = new_data.get("technologies", [])
        base_techs = merged["technologies"]
        for tech in new_techs:
            tech_id = tech.get("id")
            dup = next((t for t in base_techs if t.get("id") == tech_id), None)
            if dup:
                for k, v in tech.items():
                    if v and not dup.get(k):
                        dup[k] = v
            else:
                base_techs.append(tech)

        # 4. Слияние repositories
        new_repos = new_data.get("repositories", [])
        if new_repos:
            base_repos = merged["repositories"]
            for repo in new_repos:
                url = repo.get("url")
                if not url:
                    continue
                dup = next((r for r in base_repos if r.get("url") == url), None)
                if dup:
                    for k in ["detected_languages", "file_structure"]:
                        base_arr = dup.setdefault(k, [])
                        new_arr = repo.get(k, [])
                        if new_arr:
                            for item in new_arr:
                                if item not in base_arr:
                                    base_arr.append(item)
                else:
                    base_repos.append(repo)

        # 5. Слияние extracted_facts
        new_facts = new_data.get("extracted_facts", [])
        if new_facts:
            base_facts = merged["extracted_facts"]
            for fact in new_facts:
                fact_id = fact.get("id")
                dup = next((f for f in base_facts if f.get("id") == fact_id), None)
                if dup:
                    if dup.get("description") != fact.get("description"):
                        new_id = f"{fact_id}_conflict"
                        fact_copy = fact.copy()
                        fact_copy["id"] = new_id
                        self.logger.warning(f"Конфликт факта с ID '{fact_id}'. Сохранен дополнительный факт с ID '{new_id}'.")
                        base_facts.append(fact_copy)
                else:
                    base_facts.append(fact)

        # 6. Слияние metrics
        new_metrics = new_data.get("metrics", [])
        base_metrics = merged["metrics"]
        for met in new_metrics:
            met_id = met.get("id")
            dup = next((m for m in base_metrics if m.get("id") == met_id), None)
            if dup:
                for k, v in met.items():
                    if v and not dup.get(k):
                        dup[k] = v
            else:
                base_metrics.append(met)

        # 7. Слияние lessons_learned
        new_lessons = new_data.get("lessons_learned", [])
        if new_lessons:
            base_lessons = merged["lessons_learned"]
            for lesson in new_lessons:
                lesson_id = lesson.get("id")
                if not any(l.get("id") == lesson_id for l in base_lessons):
                    base_lessons.append(lesson)

        # 8. Слияние decisions
        new_decisions = new_data.get("decisions", [])
        base_decisions = merged["decisions"]
        for dec in new_decisions:
            dec_id = dec.get("id")
            dup = next((d for d in base_decisions if d.get("id") == dec_id), None)
            if dup:
                for k, v in dec.items():
                    if v and not dup.get(k):
                        dup[k] = v
            else:
                base_decisions.append(dec)

        # 9. Слияние evidence
        new_evidence = new_data.get("evidence", [])
        base_evidence = merged["evidence"]
        for ev in new_evidence:
            ev_id = ev.get("id")
            dup = next((e for e in base_evidence if e.get("id") == ev_id), None)
            if dup:
                for k, v in ev.items():
                    if v and not dup.get(k):
                        dup[k] = v
            else:
                base_evidence.append(ev)

        return merged

    def execute(self) -> Dict[str, Any]:
        """Запуск агента сбора контекста."""
        self.logger.info("Начало работы Context Manager Agent...")
        
        # 1. Сканируем папку input/ и парсим файлы
        combined_text, structured_json = self.scan_input_directory()
        
        llm_json = {}
        
        # 2. Если извлечен сырой текст, нормализуем его через LLM
        if combined_text:
            prompt_template = self.get_prompt_template()
            if prompt_template:
                prompt = prompt_template.replace("{{RAW_EXTRACTED_TEXT}}", combined_text)
                self.logger.info("Отправка сырого текста в LLM для структурирования...")
                try:
                    response_text = self.call_llm(prompt)
                    # Очищаем ответ от возможных Markdown-тегов ```json
                    if response_text.startswith("```json"):
                        response_text = response_text[7:]
                    if response_text.endswith("```"):
                        response_text = response_text[:-3]
                    response_text = response_text.strip()
                    
                    llm_json = json.loads(response_text)
                    self.logger.info("Данные успешно распарсены и структурированы через LLM.")
                except json.JSONDecodeError as jde:
                    self.logger.error(f"Не удалось распарсить ответ LLM как JSON: {jde}")
                    self.logger.debug(f"Ответ LLM: {response_text}")
                except Exception as e:
                    self.logger.error(f"Ошибка работы LLM-нормализатора: {e}")
            else:
                self.logger.error("Не найден шаблон промпта для context_manager.")
        else:
            self.logger.info("Сырой текст в папке input/ отсутствует.")

        # 3. Мержим данные от LLM и спарсенные JSON-файлы
        final_knowledge = self.merge_knowledge_data(structured_json, llm_json)

        # Обеспечиваем соответствие обязательным полям схемы (если они пусты)
        if "skills" not in final_knowledge:
            final_knowledge["skills"] = []
        if "repositories" not in final_knowledge:
            final_knowledge["repositories"] = []
        if "extracted_facts" not in final_knowledge:
            final_knowledge["extracted_facts"] = []
        if "lessons_learned" not in final_knowledge:
            final_knowledge["lessons_learned"] = []
        if "projects" not in final_knowledge:
            final_knowledge["projects"] = []
        if "technologies" not in final_knowledge:
            final_knowledge["technologies"] = []
        if "evidence" not in final_knowledge:
            final_knowledge["evidence"] = []
        if "decisions" not in final_knowledge:
            final_knowledge["decisions"] = []
        if "metrics" not in final_knowledge:
            final_knowledge["metrics"] = []

        # Заполняем дефолтные обязательные поля для корректного прохождения валидации и расчета графа
        for ev in final_knowledge["evidence"]:
            if isinstance(ev, dict) and ev.get("confidence") is None:
                ev_type = ev.get("type", "")
                if "code" in ev_type or "commit" in ev_type:
                    ev["confidence"] = 1.0
                elif "cert" in ev_type:
                    ev["confidence"] = 0.90
                elif "document" in ev_type or "resume" in ev_type:
                    ev["confidence"] = 0.85
                else:
                    ev["confidence"] = 0.60

        for fact in final_knowledge["extracted_facts"]:
            if isinstance(fact, dict):
                # Расчет достоверности факта по связанным доказательствам (Evidence)
                fact_ev_ids = fact.get("evidence", [])
                fact_confidences = []
                for ev_id in fact_ev_ids:
                    linked_ev = next((e for e in final_knowledge["evidence"] if e.get("id") == ev_id), None)
                    if linked_ev:
                        fact_confidences.append(linked_ev.get("confidence", 0.85))
                if fact_confidences:
                    fact["confidence"] = max(fact_confidences)
                elif fact.get("confidence") is None:
                    fact["confidence"] = 0.85 if fact.get("source") == "resume" else 0.60

        for met in final_knowledge["metrics"]:
            if isinstance(met, dict) and met.get("confidence") is None:
                linked_fact = next((f for f in final_knowledge["extracted_facts"] if f.get("id") == met.get("fact_id")), None)
                met["confidence"] = linked_fact.get("confidence", 0.85) if linked_fact else 0.85

        for lesson in final_knowledge["lessons_learned"]:
            if isinstance(lesson, dict) and lesson.get("confidence") is None:
                lesson["confidence"] = 0.85

        for repo in final_knowledge["repositories"]:
            if isinstance(repo, dict):
                if "url" in repo:
                    if "id" not in repo or not repo["id"]:
                        repo["id"] = repo["url"].rstrip("/").split("/")[-1].lower()
                    if "name" not in repo or not repo["name"]:
                        repo["name"] = repo["url"].rstrip("/").split("/")[-1]

        # Расчет верификации и достоверности навыков (Skill)
        for skill in final_knowledge["skills"]:
            if isinstance(skill, dict):
                if "category" not in skill:
                    skill["category"] = "Uncategorized"
                
                # Поиск доказательств исходного кода
                has_code_evidence = False
                for ev_id in skill.get("evidence", []):
                    linked_ev = next((e for e in final_knowledge["evidence"] if e.get("id") == ev_id), None)
                    if linked_ev and linked_ev.get("type") in ["code_file", "git_commit"]:
                        has_code_evidence = True
                        break
                skill["verification_status"] = "verified" if (has_code_evidence or skill.get("repositories")) else "unverified"
                
                # Достоверность: 0.7 * max(Fact.confidence) + 0.3 * avg(Project.confidence)
                skill_facts_conf = [
                    f.get("confidence", 0.85) for f in final_knowledge["extracted_facts"]
                    if skill.get("id") in f.get("technologies", []) or f.get("id") in skill.get("facts", [])
                ]
                skill_projs_conf = [
                    p.get("confidence", 0.85) for p in final_knowledge["projects"]
                    if skill.get("id") in p.get("technologies", []) or p.get("id") in skill.get("projects", [])
                ]
                
                max_fact_conf = max(skill_facts_conf) if skill_facts_conf else 0.5
                avg_proj_conf = sum(skill_projs_conf)/len(skill_projs_conf) if skill_projs_conf else 0.5
                
                if skill_facts_conf or skill_projs_conf:
                    skill["confidence"] = round(0.7 * max_fact_conf + 0.3 * avg_proj_conf, 2)
                elif skill.get("confidence") is None:
                    skill["confidence"] = 0.5

        # 3.5. Сохранение сырых структурированных данных резюме в storage/normalized/resume.json
        out_paths = self.paths.get("paths", {})
        raw_resume = final_knowledge.get("raw_resume_data", {})
        if raw_resume:
            normalized_dir = out_paths.get("storage_normalized_dir", "storage/normalized")
            if not os.path.isabs(normalized_dir):
                normalized_dir = os.path.join("c:/projects/portfolio_factory", normalized_dir)
            os.makedirs(normalized_dir, exist_ok=True)
            resume_file = os.path.join(normalized_dir, "resume.json")
            try:
                with open(resume_file, "w", encoding="utf-8") as f:
                    json.dump(raw_resume, f, indent=2, ensure_ascii=False)
                self.logger.info(f"Сырые данные резюме вынесены в слой normalized: {resume_file}")
            except Exception as e:
                self.logger.error(f"Не удалось сохранить normalized резюме: {e}")

        # Копируем final_knowledge для сохранения без raw_resume_data в слой знаний
        knowledge_to_save = final_knowledge.copy()
        if "raw_resume_data" in knowledge_to_save:
            del knowledge_to_save["raw_resume_data"]

        # 4. Валидируем результирующий JSON
        self.logger.info("Валидация результирующих данных по JSON-схеме...")
        is_valid = self.validate_data(knowledge_to_save, "knowledge.schema.json")
        
        if not is_valid:
            self.logger.error("Итоговый JSON не прошел валидацию схемы! Сохранение прервано.")
            raise ValueError("Итоговая структура базы знаний не валидна по схеме knowledge.schema.json")

        # 5. Записываем итоговый файл knowledge.json
        knowledge_file = out_paths.get("knowledge_file", "storage/knowledge/knowledge.json")
        if not os.path.isabs(knowledge_file):
            knowledge_file = os.path.join("c:/projects/portfolio_factory", knowledge_file)
            
        dir_name = os.path.dirname(knowledge_file)
        if dir_name and not os.path.exists(dir_name):
            os.makedirs(dir_name, exist_ok=True)
            
        try:
            with open(knowledge_file, "w", encoding="utf-8") as f:
                json.dump(knowledge_to_save, f, indent=2, ensure_ascii=False)
            self.logger.info(f"Файл базы знаний успешно сохранен: {knowledge_file}")
        except Exception as e:
            self.logger.error(f"Не удалось сохранить итоговый knowledge.json: {e}")
            raise e

        self.logger.info("Работа Context Manager Agent успешно завершена.")
        return final_knowledge

if __name__ == "__main__":
    # Локальный запуск для ручного тестирования
    agent = ContextManagerAgent()
    agent.execute()
